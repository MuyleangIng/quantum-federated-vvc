"""Generate QE-SAC-FL system explanation Word document."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import docx.oxml

doc = Document()

# Title
title = doc.add_heading('QE-SAC-FL: System Components & Features', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('Ing Muyleang — Pukyong National University, Quantum Computing Lab').alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('For team discussion — April 2026').alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('')

# ── SECTION 1: Overview ──────────────────────────────────────────────────────
doc.add_heading('1. What We Built — System Overview', 1)
p = doc.add_paragraph()
p.add_run('QE-SAC-FL').bold = True
p.add_run(' is a Federated Quantum-Enhanced Soft Actor-Critic framework for Volt-VAR Control (VVC) across multiple power utilities with different grid sizes. It extends Lin et al. (2025) QE-SAC from a single utility to a federated multi-utility setting.')

doc.add_paragraph('')
doc.add_paragraph('The system has two main innovations:')
doc.add_paragraph('1. Identifies and solves a new problem called heterogeneous FL', style='List Number')
doc.add_paragraph('2. Enables FL across heterogeneous grids using only 280 parameters per round', style='List Number')

# ── SECTION 2: Full Architecture ─────────────────────────────────────────────
doc.add_heading('2. Full System Architecture — All Components', 1)

doc.add_paragraph('The full pipeline for each client (utility) is:')

# Architecture diagram as table
table = doc.add_table(rows=1, cols=1)
table.style = 'Table Grid'
cell = table.rows[0].cells[0]
cell.text = (
    "obs [B, obs_dim]  ← different per client (42, 111, 345)\n"
    "  ↓\n"
    "[1] LocalEncoder (PRIVATE — never shared)\n"
    "    MLP: obs_dim → 64 → 32    OR    GNN: graph nodes → pool → 32\n"
    "  ↓\n"
    "  h [B, 32]\n"
    "  ↓\n"
    "[2] SharedEncoderHead (FEDERATED — 264 params)\n"
    "    Linear(32 → 8) + Tanh × π\n"
    "  ↓\n"
    "  z [B, 8]  ∈ (−π, π)   ← ALL clients same space after alignment\n"
    "  ↓\n"
    "[3] VQC — Variational Quantum Circuit (FEDERATED — 16 params)\n"
    "    8 qubits, 2 layers: RY(encoding) → CNOT → RX(trainable) → PauliZ measurement\n"
    "  ↓\n"
    "  q [B, 8]  ∈ (−1, 1)\n"
    "  ↓\n"
    "[4] Action Heads (PRIVATE — never shared)\n"
    "    N × Linear(8, |Ai|) + Softmax  (one per device type)\n"
    "  ↓\n"
    "  action: [cap1, cap2, reg_tap]"
)
cell.paragraphs[0].runs[0].font.name = 'Courier New'
cell.paragraphs[0].runs[0].font.size = Pt(9)

doc.add_paragraph('')

# ── SECTION 3: Component Details ─────────────────────────────────────────────
doc.add_heading('3. Component Details', 1)

# Component 1
doc.add_heading('Component 1 — LocalEncoder (Private)', 2)
rows = [
    ['What it does', 'Compresses grid observation into 32-dim hidden vector'],
    ['Who has it', 'Each client has its OWN — never shared with server'],
    ['Why private', 'Contains grid topology information — commercially sensitive, security risk'],
    ['Two variants', 'MLP: flat compression  |  GNN: topology-aware (uses adjacency matrix)'],
    ['Parameters', '~2,700 params (13-bus) to ~22,000 params (123-bus) — scales with grid size'],
    ['Key point', 'Because this stays private, utilities never reveal their grid topology'],
]
t = doc.add_table(rows=len(rows)+1, cols=2)
t.style = 'Table Grid'
t.rows[0].cells[0].text = 'Property'
t.rows[0].cells[1].text = 'Detail'
for i, (k, v) in enumerate(rows):
    t.rows[i+1].cells[0].text = k
    t.rows[i+1].cells[1].text = v
doc.add_paragraph('')

# Component 2
doc.add_heading('Component 2 — SharedEncoderHead (Federated)', 2)
rows = [
    ['What it does', 'Maps 32-dim local feature to 8-dim quantum input in (−π, π)'],
    ['Who has it', 'ALL clients share the same weights — this is what gets FedAvg\'d'],
    ['Parameters', '32×8 + 8 = 264 params'],
    ['Why shared', 'Forces all clients into the SAME latent geometry — solves heterogeneous FL'],
    ['Key constraint', 'Output is in (−π, π) range (via Tanh×π) — required for VQC encoding'],
    ['This is new', 'Lin et al. (2025) does NOT have this — they train one client only'],
]
t = doc.add_table(rows=len(rows)+1, cols=2)
t.style = 'Table Grid'
t.rows[0].cells[0].text = 'Property'
t.rows[0].cells[1].text = 'Detail'
for i, (k, v) in enumerate(rows):
    t.rows[i+1].cells[0].text = k
    t.rows[i+1].cells[1].text = v
doc.add_paragraph('')

# Component 3
doc.add_heading('Component 3 — VQC (Variational Quantum Circuit, Federated)', 2)
rows = [
    ['What it does', 'Processes 8-dim latent vector using quantum operations'],
    ['Circuit', '8 qubits, 2 layers: RY(z_i) → CNOT chain → RX(θ_i) → PauliZ measurement'],
    ['Parameters', '2 layers × 8 qubits = 16 trainable rotation angles'],
    ['Who has it', 'ALL clients share same VQC weights — also FedAvg\'d'],
    ['From Lin et al.', 'Yes — we keep the exact same VQC as Lin et al. (2025), unchanged'],
    ['Key property', '16 params total — much smaller than any classical MLP (min ~144 params)'],
]
t = doc.add_table(rows=len(rows)+1, cols=2)
t.style = 'Table Grid'
t.rows[0].cells[0].text = 'Property'
t.rows[0].cells[1].text = 'Detail'
for i, (k, v) in enumerate(rows):
    t.rows[i+1].cells[0].text = k
    t.rows[i+1].cells[1].text = v
doc.add_paragraph('')

# Component 4
doc.add_heading('Component 4 — Action Heads (Private)', 2)
rows = [
    ['What it does', 'Maps 8-dim VQC output to probability over each device\'s actions'],
    ['Who has it', 'Each client has its own — never shared'],
    ['Structure', 'N × Linear(8, |Ai|) + Softmax, one head per controllable device'],
    ['Example', '13-bus: 3 heads for [cap1, cap2, reg] → [2, 2, 33] actions each'],
    ['Why private', 'Different utilities have different devices — can\'t share'],
]
t = doc.add_table(rows=len(rows)+1, cols=2)
t.style = 'Table Grid'
t.rows[0].cells[0].text = 'Property'
t.rows[0].cells[1].text = 'Detail'
for i, (k, v) in enumerate(rows):
    t.rows[i+1].cells[0].text = k
    t.rows[i+1].cells[1].text = v
doc.add_paragraph('')

# Component 5
doc.add_heading('Component 5 — FedAvg Protocol (What Gets Sent Each Round)', 2)
rows = [
    ['Sent to server', 'SharedEncoderHead weights (264) + VQC weights (16) = 280 params = 1.1 KB'],
    ['NOT sent', 'LocalEncoder, Critics, Action Heads, Replay Buffer'],
    ['Aggregation', 'Simple weighted average across all clients'],
    ['Frequency', 'Once per FL round (every 1,000 local steps)'],
    ['vs Classical SAC-FL', '107,520 params per round → 383× more expensive'],
    ['O(1) property', '280 is CONSTANT regardless of grid size (42-dim or 345-dim obs)'],
]
t = doc.add_table(rows=len(rows)+1, cols=2)
t.style = 'Table Grid'
t.rows[0].cells[0].text = 'Property'
t.rows[0].cells[1].text = 'Detail'
for i, (k, v) in enumerate(rows):
    t.rows[i+1].cells[0].text = k
    t.rows[i+1].cells[1].text = v
doc.add_paragraph('')

# ── SECTION 4: The heterogeneous FL Problem ───────────────────────────────────────────────
doc.add_heading('4. The heterogeneous FL Problem — Why It Exists, Why It Matters', 1)

doc.add_heading('What is heterogeneous FL?', 2)
doc.add_paragraph(
    'heterogeneous latent space mismatch (heterogeneous FL) is a new problem that appears ONLY when '
    'you try to federate quantum RL across clients with different observation dimensions.'
)

doc.add_heading('Why it happens:', 2)
p = doc.add_paragraph()
p.add_run('Step 1: ').bold = True
p.add_run('Client A (13-bus, obs=42) trains LocalEncoder_A → produces z_A in subspace_A')
p = doc.add_paragraph()
p.add_run('Step 2: ').bold = True
p.add_run('Client B (34-bus, obs=111) trains LocalEncoder_B → produces z_B in subspace_B')
p = doc.add_paragraph()
p.add_run('Step 3: ').bold = True
p.add_run('subspace_A ≠ subspace_B (different dims, different training data)')
p = doc.add_paragraph()
p.add_run('Step 4: ').bold = True
p.add_run('FedAvg averages VQC gradients: ∂L_A/∂θ ≈ −∂L_B/∂θ → average ≈ 0')
p = doc.add_paragraph()
p.add_run('Result: ').bold = True
p.add_run('VQC gradient norm collapses to ~10⁻³ (barren plateau) → learning stops')

doc.add_heading('Why no prior work solves this:', 2)
doc.add_paragraph(
    'All existing quantum FL papers use HOMOGENEOUS clients (same obs_dim, same grid). '
    'heterogeneous FL only appears in HETEROGENEOUS settings (different grid sizes). '
    'This is the real-world scenario for multi-utility VVC — '
    'each utility has a different sized grid.'
)

doc.add_heading('Our fix — AlignedCAE:', 2)
doc.add_paragraph(
    'Force ALL clients to use the SAME SharedEncoderHead (W_shared). '
    'Now z_A = W_shared × h_A and z_B = W_shared × h_B → both in the SAME subspace. '
    'Gradients from A and B now point in the same direction → FedAvg works → VQC learns.'
)

# ── SECTION 5: Features Summary ──────────────────────────────────────────────
doc.add_heading('5. Full Feature List', 1)

features = [
    ('Multi-utility FL', 'Train 3 utilities simultaneously, share knowledge without sharing data'),
    ('Heterogeneous grids', 'Works with 13-bus (obs=42), 34-bus (obs=111), 123-bus (obs=345)'),
    ('heterogeneous FL solved', 'AlignedCAE forces shared latent geometry — first work to solve this'),
    ('Barren plateau prevention', 'Aligned gradients prevent VQC from collapsing during FL'),
    ('Privacy-preserving', 'Grid topology (LocalEncoder) never leaves the client'),
    ('O(1) communication', '280 params/round regardless of grid size — constant cost'),
    ('383× communication reduction', 'vs classical SAC-FL (107,520 params/round)'),
    ('Quantum circuit preserved', 'Same VQC as Lin et al. (2025) — fair comparison'),
    ('GNN variant', 'Optional GNN LocalEncoder for topology-aware encoding (fault recovery)'),
    ('Personalized FL', 'After federated pre-training, fine-tune per client → further improvement'),
    ('Statistical validation', 'n=5 seeds, Cohen\'s d, p-value, Bonferroni correction'),
    ('Real OpenDSS validation', '3-phase AC physics, real IEEE feeder data'),
]

t = doc.add_table(rows=len(features)+1, cols=2)
t.style = 'Table Grid'
t.rows[0].cells[0].text = 'Feature'
t.rows[0].cells[1].text = 'Description'
for i, (f, d_) in enumerate(features):
    t.rows[i+1].cells[0].text = f
    t.rows[i+1].cells[1].text = d_

doc.add_paragraph('')

# ── SECTION 6: What is NOT our contribution ───────────────────────────────────
doc.add_heading('6. What We Did NOT Invent (Be Honest)', 1)
not_ours = [
    ('VQC circuit design', 'From Lin et al. (2025) — we use the same 8-qubit, 2-layer circuit'),
    ('SAC algorithm', 'Standard Soft Actor-Critic (Haarnoja et al. 2018)'),
    ('FedAvg', 'Standard federated averaging (McMahan et al. 2017)'),
    ('IEEE test feeders', 'Standard benchmarks (IEEE PES 1992)'),
    ('CAE (autoencoder)', 'Concept from Lin et al. — we extend it to the shared/private split'),
]
t = doc.add_table(rows=len(not_ours)+1, cols=2)
t.style = 'Table Grid'
t.rows[0].cells[0].text = 'Component'
t.rows[0].cells[1].text = 'Source'
for i, (f, d_) in enumerate(not_ours):
    t.rows[i+1].cells[0].text = f
    t.rows[i+1].cells[1].text = d_

# ── SECTION 7: Open Questions ─────────────────────────────────────────────────
doc.add_heading('7. Open Questions for Wednesday Discussion', 1)
questions = [
    '1. Is the variance reduction (±0.016 vs ±0.055) sufficient as a main result, or do we need larger reward improvement?',
    '2. Should we add batteries to match Lin et al. exactly, or is the simplified action space acceptable?',
    '3. Is the GNN variant worth including, or focus on MLP only?',
    '4. Target venue: IEEE Transactions on Smart Grid — is this the right fit?',
    '5. Are there additional baselines we should compare against?',
]
for q in questions:
    doc.add_paragraph(q, style='List Bullet')

# Save
out = 'artifacts/QE_SAC_FL_System_Explanation.docx'
doc.save(out)
print(f'Saved → {out}')
